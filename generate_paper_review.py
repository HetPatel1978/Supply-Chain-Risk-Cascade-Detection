"""Generate a black-and-white Word document for the two-paper assignment review."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0, 0, 0)
WHITE = RGBColor(255, 255, 255)
LIGHT_GREY = RGBColor(220, 220, 220)
MID_GREY = RGBColor(150, 150, 150)


def set_cell_bg(cell, hex_color: str):
    """Fill a table cell background."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def cell_color_rgb(cell, r, g, b):
    hex_color = f"{r:02X}{g:02X}{b:02X}"
    set_cell_bg(cell, hex_color)


def set_borders(table):
    """Add thin black borders to every cell."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def para(doc, text, bold=False, size=11, color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_before=0, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = BLACK
    # bottom border via paragraph border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = BLACK
    return p


def section_label(doc, number, title):
    """Numbered section label e.g. '1. What is the aim?'"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{number}.  {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    return p


def bullet(doc, text, indent=0.3):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK
    return p


def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_kv_table(doc, rows, col_widths=(1.8, 4.4)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_borders(table)
    for i, (key, val) in enumerate(rows):
        table.columns[0].width = Inches(col_widths[0])
        table.columns[1].width = Inches(col_widths[1])
        c0 = table.cell(i, 0)
        c1 = table.cell(i, 1)
        if i % 2 == 0:
            cell_color_rgb(c0, 230, 230, 230)
            cell_color_rgb(c1, 230, 230, 230)
        else:
            cell_color_rgb(c0, 245, 245, 245)
            cell_color_rgb(c1, 245, 245, 245)
        r0 = c0.paragraphs[0].add_run(key)
        r0.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = BLACK
        r1 = c1.paragraphs[0].add_run(val)
        r1.font.size = Pt(10)
        r1.font.color.rgb = BLACK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_comparison_table(doc):
    headers = ["", "REBEL (EMNLP 2021)", "RAG (NeurIPS 2020)"]
    rows = [
        ("Pipeline step", "Steps 1-3: Relation Extraction", "Step 5: Multi-hop RAG"),
        ("Input", "Raw sentence", "Query + document index"),
        ("Output", "(head, relation, tail) triples", "Natural language answer"),
        ("Core idea", "RE as seq2seq generation", "Retrieval + generation jointly"),
        ("System analogy", "RoBERTa fine-tune + Groq extractor", "Groq LLM over graph-retrieved context"),
        ("Key lesson", "Joint NER+RE avoids spaCy entity-detection failures",
         "Dense retrieval + modular index enables live updates"),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_borders(table)
    # header row
    widths = (1.5, 2.5, 2.6)
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        cell_color_rgb(c, 30, 30, 30)
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = WHITE
        table.columns[j].width = Inches(widths[j])
    # data rows
    for i, (col0, col1, col2) in enumerate(rows):
        shade = (220, 220, 220) if i % 2 == 0 else (245, 245, 245)
        for j, text in enumerate((col0, col1, col2)):
            c = table.cell(i + 1, j)
            cell_color_rgb(c, *shade)
            r = c.paragraphs[0].add_run(text)
            r.font.size = Pt(10)
            r.font.color.rgb = BLACK
            if j == 0:
                r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_challenge_table(doc, rows):
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_borders(table)
    for j, h in enumerate(("Challenge", "Details")):
        c = table.cell(0, j)
        cell_color_rgb(c, 30, 30, 30)
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = WHITE
        table.columns[j].width = Inches(2.0 if j == 0 else 4.2)
    for i, (ch, dt) in enumerate(rows):
        shade = (220, 220, 220) if i % 2 == 0 else (245, 245, 245)
        for j, text in enumerate((ch, dt)):
            c = table.cell(i + 1, j)
            cell_color_rgb(c, *shade)
            r = c.paragraphs[0].add_run(text)
            r.font.size = Pt(10)
            r.font.color.rgb = BLACK
            if j == 0:
                r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── Cover / Title block ───────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("PAPER REVIEW ASSIGNMENT")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = BLACK

    para(doc, "Supply Chain Risk Detection — NLP Pipeline",
         bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "Two Selected Papers (excludes Kaur & AlMahri)",
         size=10, color=RGBColor(80, 80, 80), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

    divider(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # PAPER 1: REBEL
    # ═══════════════════════════════════════════════════════════════════════════
    heading1(doc, "Paper 1 — REBEL: Relation Extraction By End-to-end Language Generation")

    add_kv_table(doc, [
        ("Authors",   "Pere-Lluís Cabot, Roberto Navigli"),
        ("Venue",     "Findings of EMNLP 2021"),
        ("Access",    "ACL Anthology — aclanthology.org/2021.findings-emnlp.204"),
        ("Backbone",  "BART-large (seq2seq Transformer)"),
        ("Task",      "Joint Named Entity Recognition + Relation Extraction"),
    ])

    # 1. Aim
    section_label(doc, 1, "What is the Aim of the Research?")
    para(doc,
         "REBEL aims to perform relation extraction (RE) as a sequence-to-sequence "
         "generation task. Instead of the traditional two-stage pipeline — first detect "
         "entity spans with a NER model, then classify the relation between each pair — "
         "REBEL trains a single BART-based model that reads a raw sentence and generates "
         "(subject, relation, object) triplets directly as a linearized token sequence. "
         "The overarching goal is a unified, end-to-end RE system that generalizes across "
         "many domains and relation types without requiring task-specific entity detectors.",
         size=10.5, space_after=4)

    # 2. Challenges
    section_label(doc, 2, "What Are the Scientific Challenges?")
    add_challenge_table(doc, [
        ("Pipeline error cascading",
         "Separate NER → RE models compound mistakes; an entity missed by NER can never be extracted."),
        ("Relation type explosion",
         "Hundreds of relation types exist across corpora; learning one model to cover all is hard."),
        ("Cross-sentence extraction",
         "Relevant entity pairs often span multiple sentences; standard token-pair classifiers ignore this."),
        ("Structured output from generation",
         "Teaching seq2seq to produce well-formed triples requires special output linearization."),
        ("Noisy distant supervision",
         "Training data is auto-generated by aligning Wikipedia text with Wikidata — label noise at scale."),
    ])

    # 3. Contribution
    section_label(doc, 3, "Describe the Contribution of This Research")
    for b in [
        "REBEL linearization: introduces special delimiter tokens (<triplet>, <subj>, <obj>) so BART "
        "generates structured triples in a single forward pass — no separate span detection needed.",
        "REBEL dataset: a large automatically-constructed RE corpus covering 220 relation types "
        "across all of English Wikipedia, used for pre-training.",
        "State-of-the-art results on DocRED, NYT, and CoNLL04 benchmarks without task-specific NER, "
        "and with competitive zero-shot transfer to new domains.",
        "Overlapping triplets (one entity in multiple relations) are handled naturally by generation, "
        "which discriminative span-pair models struggle with.",
    ]:
        bullet(doc, b)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 4. Differences from prior work
    section_label(doc, 4, "How Does This Differ from Existing Related Work?")
    add_kv_table(doc, [
        ("Pipeline models (SpERT, ATLOP)",
         "Require gold/predicted entity spans before RE; error from NER propagates. REBEL does NER+RE jointly."),
        ("Span-pair classifiers (CasRel, TPLinker)",
         "Tag every token-pair combination — quadratic complexity. REBEL is linear in output length."),
        ("Earlier seq2seq RE (TANL — Paolini et al.)",
         "Also generative but requires structured augmentation markup; REBEL is simpler and pre-trained at larger scale."),
        ("Supervised domain-specific RE",
         "Require domain-specific training sets. REBEL's Wikipedia pre-training allows zero-shot transfer."),
    ])

    # 5. Strengths / Weaknesses
    section_label(doc, 5, "Strengths and Weaknesses")
    para(doc, "Strengths:", bold=True, size=10.5, space_after=1)
    for b in [
        "End-to-end training eliminates cascading pipeline errors from separate NER.",
        "Handles overlapping and nested entity mentions naturally.",
        "Pre-training on 220 relations enables strong transfer to new domains.",
        "Simple unified architecture with no task-specific components.",
    ]:
        bullet(doc, b)
    para(doc, "Weaknesses:", bold=True, size=10.5, space_before=4, space_after=1)
    for b in [
        "Autoregressive generation is slower at inference than discriminative classifiers.",
        "Can hallucinate entity spans not present in the input (a known seq2seq failure mode).",
        "Relation types outside the pre-training distribution are hard to generalize to.",
        "Label noise from distant supervision in pre-training data can hurt precision.",
    ]:
        bullet(doc, b)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 6. Relevance
    section_label(doc, 6, "Relevance to the Project & Lessons Learned")
    para(doc,
         "REBEL's task is exactly what Steps 1–3 of the supply chain pipeline do — read a 10-K "
         "sentence, identify company entities, and output (head, relation, tail) triples over a "
         "fixed vocabulary. The rule-based extractor and Groq zero-shot extractor both approximate "
         "what REBEL does in a trained end-to-end manner.", size=10.5, space_after=4)
    para(doc, "Key lessons:", bold=True, size=10.5, space_after=1)
    for b in [
        "Linearization format matters for LLM prompting: REBEL's delimiters enforce structure, "
        "which inspired the Groq prompt's strict JSON schema with explicit head/tail/relation fields.",
        "Joint NER+RE avoids the entity-detection bottleneck hit when spaCy's en_core_web_sm "
        "failed to tag 'Arm Limited' as ORG — a generative model would have handled this naturally.",
        "Fine-tuning on REFinD (Step 1) is the right direction: REBEL shows pre-training on a "
        "large RE corpus then fine-tuning on domain data achieves SOTA.",
        "REBEL could serve as a stronger fine-tuned baseline to compare against the rule-based "
        "and zero-shot approaches in the final evaluation.",
    ]:
        bullet(doc, b)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # PAPER 2: RAG
    # ═══════════════════════════════════════════════════════════════════════════
    heading1(doc, "Paper 2 — Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks")

    add_kv_table(doc, [
        ("Authors",   "Lewis, Perez, Piktus, Petroni, Karpukhin, Goyal, Küttler, Lewis, Yih, "
                      "Rocktäschel, Riedel, Kiela"),
        ("Venue",     "Advances in Neural Information Processing Systems (NeurIPS 2020)"),
        ("Access",    "proceedings.neurips.cc/paper/2020 — hash 6b493230205f780e1bc26945df7481e5"),
        ("Backbone",  "DPR retriever + BART-large generator"),
        ("Task",      "Open-domain QA, fact verification, knowledge-intensive NLP generation"),
    ])

    # 1. Aim
    section_label(doc, 1, "What is the Aim of the Research?")
    para(doc,
         "RAG aims to give large language models updatable, external memory by coupling them "
         "with a retrieval component over a document index. The model retrieves the most relevant "
         "passages for a query at inference time, then conditions the generator on both the query "
         "and the retrieved evidence to produce an answer. The goal is to improve performance on "
         "knowledge-intensive NLP tasks while making the model's knowledge transparent and "
         "modifiable without full retraining.",
         size=10.5, space_after=4)

    # 2. Challenges
    section_label(doc, 2, "What Are the Scientific Challenges?")
    add_challenge_table(doc, [
        ("Knowledge staleness",
         "LLM weights encode knowledge frozen at training time; new facts require expensive retraining."),
        ("Hallucination",
         "Parametric-only generation invents facts not in training data, especially for rare/recent entities."),
        ("Sparse vs dense retrieval",
         "Classic BM25 misses semantically relevant passages that use different vocabulary."),
        ("End-to-end training",
         "Integrating discrete document retrieval with differentiable generation requires marginalizing over docs."),
        ("Multi-document reasoning",
         "An answer may require combining evidence from several retrieved passages, not just one."),
    ])

    # 3. Contribution
    section_label(doc, 3, "Describe the Contribution of This Research")
    for b in [
        "RAG architecture: a Dense Passage Retriever (DPR) selects the top-k documents for a query; "
        "a BART seq2seq generator conditions on the query + retrieved context to produce output.",
        "RAG-Sequence variant: uses the same retrieved document for every output token.",
        "RAG-Token variant: can draw on different documents for different output tokens — richer "
        "multi-document synthesis.",
        "End-to-end training by marginalizing the generator's likelihood over top-k retrieved documents, "
        "making retrieval implicitly differentiable.",
        "State-of-the-art on Natural Questions, TriviaQA, WebQuestions, MS-MARCO, and FEVER at publication.",
        "Shows that the retrieval index can be swapped or updated post-training, giving the model "
        "dynamic knowledge without full retraining.",
    ]:
        bullet(doc, b)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 4. Differences from prior work
    section_label(doc, 4, "How Does This Differ from Existing Related Work?")
    add_kv_table(doc, [
        ("Closed-book QA (T5, GPT-3)",
         "Generate answers purely from weights — no external lookup, prone to hallucination. "
         "RAG grounds generation in retrieved evidence."),
        ("REALM (Guu et al. 2020)",
         "Also retrieval-augmented, but designed for masked-LM pre-training, not seq2seq generation."),
        ("Fusion-in-Decoder (FiD)",
         "Encodes all retrieved passages jointly, but lacks token-level document marginalization."),
        ("Open-book QA with BM25",
         "Sparse keyword retrieval misses semantic relevance. RAG uses dense embeddings via DPR."),
    ])

    # 5. Strengths / Weaknesses
    section_label(doc, 5, "Strengths and Weaknesses")
    para(doc, "Strengths:", bold=True, size=10.5, space_after=1)
    for b in [
        "Modular and updatable: swap the document index with new SEC filings without retraining the generator.",
        "Transparent retrieval: you can inspect which passages were retrieved for any answer.",
        "Strong empirical results across heterogeneous knowledge-intensive tasks.",
        "Reduces hallucination compared to parametric-only generation by anchoring answers in retrieved text.",
    ]:
        bullet(doc, b)
    para(doc, "Weaknesses:", bold=True, size=10.5, space_before=4, space_after=1)
    for b in [
        "Single-hop retrieval: the standard RAG architecture retrieves once per query; multi-hop "
        "reasoning requires extensions like iterative retrieval.",
        "Retrieval latency: dense retrieval with FAISS adds inference-time cost at scale.",
        "Index granularity matters: retrieval over whole documents performs worse than sub-paragraphs; "
        "10-K sections need careful chunking.",
        "Generator-retriever gap: if distributions differ, retrieval quality degrades.",
    ]:
        bullet(doc, b)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 6. Relevance
    section_label(doc, 6, "Relevance to the Project & Lessons Learned")
    para(doc,
         "Step 5 of the supply chain pipeline is explicitly multi-hop retrieval-augmented generation "
         "over the knowledge graph. RAG is the foundational architecture for this step. The system "
         "extends it: instead of retrieving from a flat document corpus, it traverses the NetworkX "
         "supply-chain graph to gather multi-hop context, then feeds it to the LLM for cascade-risk "
         "generation.", size=10.5, space_after=4)
    para(doc, "Key lessons:", bold=True, size=10.5, space_after=1)
    for b in [
        "Dense embeddings (sentence-transformers) beat BM25 for semantic queries — requirements.txt "
        "already includes sentence-transformers and faiss-cpu for exactly this reason.",
        "The index is decoupled from the generator: when a new 10-K filing is added, re-embed and "
        "add to FAISS without retraining the LLM.",
        "Multi-hop requires iterative retrieval: start at a node (e.g., NVIDIA), follow supplier_of "
        "edges to TSMC, then to ASML, accumulating context at each hop before querying the LLM.",
        "Passage chunking is critical: the paper shows 100-token windows significantly outperform "
        "document-level retrieval — split 10-K Item 1A into paragraph-level chunks before indexing.",
        "For evaluation (Step 6): the paper uses faithfulness (does the answer match retrieved passages?) "
        "as an intrinsic metric — maps directly to the extrinsic RAG evaluation design.",
    ]:
        bullet(doc, b)

    divider(doc)

    # ── Comparison table ──────────────────────────────────────────────────────
    heading2(doc, "Quick Comparison: Both Papers")
    add_comparison_table(doc)

    # ── Footer note ───────────────────────────────────────────────────────────
    para(doc,
         "Both papers are publicly available via ACL Anthology and NeurIPS proceedings respectively, "
         "and are also accessible on Semantic Scholar and Google Scholar. Neither paper is authored "
         "by Kaur or AlMahri.",
         size=9, color=RGBColor(100, 100, 100), space_before=6, space_after=0)

    out = "Paper_Review_Assignment.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build()
